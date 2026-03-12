"""
DOCX → Text converter.

Reads paragraphs and tables from a .docx file, yielding text in chunks
to keep memory usage low.
"""

import logging
from typing import Generator

from docx import Document

logger = logging.getLogger(__name__)

CHUNK_PARAGRAPHS = 50  # yield every N paragraphs


def convert_to_text(file_path: str) -> Generator[str, None, None]:
    """Yield text chunks from a DOCX file."""
    logger.info("DOCX convert: %s", file_path)
    doc = Document(file_path)
    buffer: list[str] = []

    # ── Paragraphs ───────────────────────────────────────────────────────
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            buffer.append(text)
        if len(buffer) >= CHUNK_PARAGRAPHS:
            yield "\n".join(buffer) + "\n"
            buffer.clear()

    # ── Tables ───────────────────────────────────────────────────────────
    for table_idx, table in enumerate(doc.tables):
        rows_text: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        if rows_text:
            buffer.append(f"[TABLE {table_idx + 1}]\n" + "\n".join(rows_text))

    if buffer:
        yield "\n".join(buffer) + "\n"
